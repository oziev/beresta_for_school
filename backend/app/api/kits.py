from __future__ import annotations

import io
import json
import logging
import uuid
from typing import Annotated, Any

import httpx
from fastapi import APIRouter, Depends, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import Response
from pydantic import ValidationError
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.config import get_settings
from app.database import get_session
from app.export.pdf import render_kit_pdf
from app.llm.cache import cache_get, cache_key_for_prompt, cache_set
from app.llm.client import GigaChatClient, GigaChatError
from app.llm.jsonutil import JSON_ONLY_SUFFIX, extract_json_object
from app.llm.prompts import (
    prompt_example_with_mistakes,
    prompt_generate_kit_from_plan,
    prompt_hot_mistakes,
    prompt_parse_lesson_plan,
)
from app.models import DesignSettings, Kit, KitItem, LessonPlan, TeacherProfile, User
from app.schemas import (
    ContentLevels,
    DesignSettingsPayload,
    DefaultPlanRequest,
    ExampleMistakesPayload,
    GenerateFromPlanRequest,
    GenerateFromPlanResponse,
    GenerateFromScratchRequest,
    KitExportRequest,
    KitItemCreate,
    KitItemRead,
    KitItemReorderRequest,
    KitItemUpdate,
    KitRead,
    KitTemplateRequest,
    LessonFeedback,
    LessonPlanRead,
    LessonStage,
    PedagogicalFeatures,
    UploadLessonPlanResponse,
)

log = logging.getLogger(__name__)
router = APIRouter(tags=["kits"])
DEMO_EMAIL = "demo@beresta.local"

LESSON_TYPES: dict[str, dict[str, Any]] = {
    "discovery": {
        "title": "Открытие нового знания",
        "icon": "💡",
        "stages": [
            ("Актуализация знаний", 5, "quiz", "cards"),
            ("Проблемная ситуация", 7, "practice", "schema"),
            ("Открытие нового знания", 12, "explanation", "memo"),
            ("Первичное закрепление", 10, "practice", "worksheet"),
            ("Самостоятельная работа", 10, "practice", "worksheet"),
        ],
    },
    "reflection": {
        "title": "Рефлексия",
        "icon": "🔄",
        "stages": [
            ("Анализ ошибок", 8, "practice", "table"),
            ("Индивидуальная работа", 15, "practice", "worksheet"),
            ("Взаимопроверка", 10, "practice", "cards"),
            ("Рефлексия", 7, "reflection", "reflection"),
        ],
    },
    "control": {
        "title": "Контроль знаний",
        "icon": "📝",
        "stages": [
            ("Организационный момент", 2, "explanation", "memo"),
            ("Инструктаж", 3, "explanation", "memo"),
            ("Контрольная работа", 30, "practice", "worksheet"),
            ("Самопроверка", 5, "practice", "table"),
            ("Анализ ошибок", 5, "reflection", "reflection"),
        ],
    },
    "combined": {
        "title": "Комбинированный ФГОС",
        "icon": "🎯",
        "stages": [
            ("Организационный момент", 2, "explanation", "memo"),
            ("Актуализация знаний", 5, "quiz", "cards"),
            ("Целеполагание", 3, "reflection", "schema"),
            ("Открытие нового знания", 12, "explanation", "memo"),
            ("Первичное закрепление", 10, "practice", "worksheet"),
            ("Самостоятельная работа", 10, "practice", "worksheet"),
            ("Рефлексия", 3, "reflection", "reflection"),
        ],
    },
}


def _default_stages(lesson_type: str) -> list[LessonStage]:
    spec = LESSON_TYPES.get(lesson_type, LESSON_TYPES["combined"])
    return [
        LessonStage(
            name=name,
            default_duration=duration,
            time_minutes=duration,
            activity_type=activity,
            needs_handout=activity != "explanation" or handout == "memo",
            recommended_handout_type=handout,
        )
        for name, duration, activity, handout in spec["stages"]
    ]


def get_llm(request: Request) -> GigaChatClient:
    return request.app.state.llm


def get_redis(request: Request):
    return getattr(request.app.state, "redis", None)


def _default_stage_from_line(line: str, idx: int) -> LessonStage:
    low = line.lower()
    if "актуал" in low or "опрос" in low:
        return LessonStage(name=line[:100], activity_type="quiz", recommended_handout_type="worksheet")
    if "нов" in low or "объяс" in low or "изуч" in low:
        return LessonStage(name=line[:100], activity_type="explanation", recommended_handout_type="memo")
    if "рефлекс" in low or "итог" in low:
        return LessonStage(name=line[:100], activity_type="reflection", recommended_handout_type="reflection")
    if "домаш" in low:
        return LessonStage(name=line[:100], activity_type="homework", recommended_handout_type="worksheet")
    return LessonStage(name=line[:100] or f"Этап {idx + 1}", activity_type="practice", recommended_handout_type="worksheet")


def _fallback_parse_stages(text: str) -> list[LessonStage]:
    lines = [ln.strip(" -—\t") for ln in text.splitlines() if len(ln.strip()) > 3]
    candidates = [ln for ln in lines if any(word in ln.lower() for word in ("этап", "актуал", "изуч", "закреп", "рефлекс", "домаш"))]
    source = candidates[:5] or lines[:5] or ["Актуализация знаний", "Изучение нового", "Закрепление", "Рефлексия"]
    return [_default_stage_from_line(line, idx) for idx, line in enumerate(source[:5])]


VALID_HANDOUT_TYPES = {"cards", "worksheet", "memo", "reflection", "table", "schema", "homework"}


def _normalize_items(
    data: dict[str, Any],
    stages: list[LessonStage],
    *,
    global_complexity: int = 2,
) -> list[dict[str, Any]]:
    """Парсит ответ LLM в нормализованные элементы комплекта.

    Если LLM не вернул items — генерируем заглушки по этапам (но всегда
    с реальными типами раздаток из плана).
    """
    raw_items = data.get("items")
    enabled_stages = [s for s in stages if s.needs_handout]
    if not isinstance(raw_items, list) or not raw_items:
        raw_items = []
        for idx, stage in enumerate(enabled_stages[:6]):
            handout = stage.recommended_handout_type if stage.recommended_handout_type in VALID_HANDOUT_TYPES else "worksheet"
            raw_items.append(
                {
                    "stage_name": stage.name,
                    "type": handout,
                    "title": f"Раздатка: {stage.name}",
                    "content_levels": {
                        "basic": f"Базовый материал для этапа «{stage.name}» (заглушка LLM).",
                        "medium": f"Средний материал для этапа «{stage.name}» (заглушка LLM).",
                        "advanced": f"Продвинутый материал для этапа «{stage.name}» (заглушка LLM).",
                    },
                }
            )
    normalized: list[dict[str, Any]] = []
    for idx, item in enumerate(raw_items[:6]):
        levels = item.get("content_levels") if isinstance(item, dict) else None
        if not isinstance(levels, dict):
            levels = {}
        content_levels = ContentLevels(
            basic=str(levels.get("basic") or "Базовый вариант материала"),
            medium=str(levels.get("medium") or "Средний вариант материала"),
            advanced=str(levels.get("advanced") or "Продвинутый вариант материала"),
        )
        item_type = (item.get("type") or "worksheet").strip().lower()
        if item_type not in VALID_HANDOUT_TYPES:
            item_type = "worksheet"
        answer_key = item.get("answer_key") if isinstance(item.get("answer_key"), dict) else None
        # Собираем доп. поля в example_mistakes (там у нас единое поле для всех "приписок")
        extras: list[dict[str, Any]] = []
        raw_mistakes = item.get("example_mistakes")
        if isinstance(raw_mistakes, list):
            extras.extend(x for x in raw_mistakes if isinstance(x, dict))
        scaffolding = item.get("scaffolding_steps")
        if isinstance(scaffolding, list) and scaffolding:
            extras.append({"kind": "scaffolding", "steps": [str(s) for s in scaffolding][:8]})
        sources = item.get("sources")
        if isinstance(sources, list) and sources:
            extras.append({"kind": "sources", "items": [s for s in sources if isinstance(s, dict)][:5]})
        normalized.append(
            {
                "stage_name": item.get("stage_name") or f"Этап {idx + 1}",
                "type": item_type,
                "title": item.get("title") or f"Раздатка {idx + 1}",
                "content_levels": content_levels.model_dump(mode="json"),
                "complexity_level": max(1, min(3, int(item.get("complexity_level") or global_complexity))),
                "complexity_distribution": item.get("complexity_distribution") or "uniform",
                "sort_order": idx,
                "teacher_notes": item.get("teacher_notes"),
                "answer_key": answer_key,
                "example_mistakes": extras or None,
            }
        )
    return normalized


async def _load_teacher_profile(db: AsyncSession) -> dict[str, Any] | None:
    """Подгружает профиль виртуального двойника учителя, если он активен."""
    res = await db.execute(select(TeacherProfile).where(TeacherProfile.user_email == DEMO_EMAIL))
    tp = res.scalar_one_or_none()
    if tp is None or not tp.enabled or not tp.profile:
        return None
    return tp.profile


async def _extract_upload_text(file: UploadFile) -> str:
    raw = await file.read()
    filename = (file.filename or "lesson.txt").lower()
    if filename.endswith(".txt"):
        return raw.decode("utf-8", errors="ignore")
    if filename.endswith(".pdf"):
        try:
            from PyPDF2 import PdfReader
        except ImportError as exc:
            raise HTTPException(status_code=500, detail="pypdf2_not_installed") from exc
        reader = PdfReader(io.BytesIO(raw))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if filename.endswith(".docx"):
        try:
            from docx import Document
        except ImportError as exc:
            raise HTTPException(status_code=500, detail="python_docx_not_installed") from exc
        doc = Document(io.BytesIO(raw))
        return "\n".join(p.text for p in doc.paragraphs)
    raise HTTPException(status_code=400, detail="unsupported_file_type")


async def _parse_stages_with_llm(text: str, llm: GigaChatClient, redis: Any) -> list[LessonStage]:
    prompt = prompt_parse_lesson_plan(text=text[:12000])
    key = cache_key_for_prompt(prompt)
    raw = await cache_get(redis, key)
    if raw is None:
        try:
            raw = await llm.chat_completion(prompt)
        except (GigaChatError, httpx.HTTPError) as exc:
            log.error("Lesson plan parsing failed: %s", exc)
            return _fallback_parse_stages(text)
        await cache_set(redis, key, raw, get_settings().llm_cache_ttl_sec)
    try:
        data = extract_json_object(raw)
        stages = [LessonStage.model_validate(x) for x in data.get("stages", [])]
        return stages or _fallback_parse_stages(text)
    except (ValueError, ValidationError, json.JSONDecodeError) as exc:
        log.error("Invalid lesson plan JSON: %s", exc)
        return _fallback_parse_stages(text)


async def _get_kit_or_404(db: AsyncSession, kit_id: uuid.UUID) -> Kit:
    res = await db.execute(select(Kit).options(selectinload(Kit.items)).where(Kit.id == kit_id))
    kit = res.scalar_one_or_none()
    if kit is None:
        raise HTTPException(status_code=404, detail="kit_not_found")
    return kit


@router.get("/lesson-types")
async def lesson_types() -> dict[str, Any]:
    return {
        key: {
            "id": key,
            "title": spec["title"],
            "icon": spec["icon"],
            "stages": [stage.model_dump(mode="json") for stage in _default_stages(key)],
        }
        for key, spec in LESSON_TYPES.items()
    }


@router.post("/generate-default-plan")
async def generate_default_plan(body: DefaultPlanRequest) -> dict[str, Any]:
    stages = _default_stages(body.lesson_type)
    return {
        "topic": body.topic,
        "grade": body.grade,
        "subject": body.subject,
        "lesson_type": body.lesson_type,
        "stages": [stage.model_dump(mode="json") for stage in stages],
    }


@router.post("/upload-lesson-plan", response_model=UploadLessonPlanResponse)
async def upload_lesson_plan(
    request: Request,
    db: Annotated[AsyncSession, Depends(get_session)],
    llm: Annotated[GigaChatClient, Depends(get_llm)],
    file: Annotated[UploadFile, File()],
    save: Annotated[bool, Form()] = False,
    name: Annotated[str | None, Form()] = None,
) -> UploadLessonPlanResponse:
    content = (await _extract_upload_text(file)).strip()
    if not content:
        raise HTTPException(status_code=400, detail="empty_lesson_plan")
    stages = await _parse_stages_with_llm(content, llm, get_redis(request))
    lesson_plan_id = None
    if save:
        plan = LessonPlan(
            user_email=None,
            name=name or file.filename or "План урока",
            original_filename=file.filename,
            content=content,
            stages=[s.model_dump(mode="json") for s in stages],
        )
        db.add(plan)
        await db.commit()
        await db.refresh(plan)
        lesson_plan_id = plan.id
    return UploadLessonPlanResponse(lesson_plan_id=lesson_plan_id, content=content, stages=stages)


@router.post("/generate-from-plan", response_model=GenerateFromPlanResponse)
async def generate_from_plan(
    request: Request,
    body: GenerateFromPlanRequest,
    db: Annotated[AsyncSession, Depends(get_session)],
    llm: Annotated[GigaChatClient, Depends(get_llm)],
) -> GenerateFromPlanResponse:
    plan: LessonPlan | None = None
    content = body.lesson_plan_content
    stages = body.stages
    if body.lesson_plan_id is not None:
        res = await db.execute(select(LessonPlan).where(LessonPlan.id == body.lesson_plan_id))
        plan = res.scalar_one_or_none()
        if plan is None:
            raise HTTPException(status_code=404, detail="lesson_plan_not_found")
        content = plan.content
        stages = [LessonStage.model_validate(x) for x in plan.stages]
    if not content and not stages:
        raise HTTPException(status_code=400, detail="lesson_plan_required")
    if stages is None:
        stages = await _parse_stages_with_llm(content or "", llm, get_redis(request))
    if plan is None and body.save_plan and content:
        plan = LessonPlan(
            user_email=None,
            name=body.plan_name or body.topic or "План урока",
            original_filename=None,
            content=content,
            stages=[s.model_dump(mode="json") for s in stages],
        )
        db.add(plan)
        await db.flush()
    # US-04: учитываем total lesson time → при необходимости отрезаем лишние этапы.
    if body.pedagogical and body.pedagogical.timing and body.pedagogical.timing.enabled:
        budget = max(5, int(body.pedagogical.timing.lesson_duration))
        used = 0
        trimmed: list[LessonStage] = []
        for s in stages:
            cost = int(s.time_minutes or s.default_duration or 5)
            if used + cost > budget and trimmed:
                break
            trimmed.append(s)
            used += cost
        if trimmed:
            stages = trimmed
    teacher_profile = await _load_teacher_profile(db)
    lesson_plan_json = json.dumps([s.model_dump(mode="json") for s in stages], ensure_ascii=False)
    prompt = prompt_generate_kit_from_plan(
        lesson_plan=lesson_plan_json,
        topic=body.topic,
        grade=body.grade,
        subject=body.subject,
        lesson_type=body.lesson_type,
        global_complexity=body.global_complexity,
        pedagogical=body.pedagogical,
        teacher_profile=teacher_profile,
    )
    redis = get_redis(request)
    key = cache_key_for_prompt(prompt)
    raw = await cache_get(redis, key)
    if raw is None:
        try:
            raw = await llm.chat_completion(prompt)
        except (GigaChatError, httpx.HTTPError) as exc:
            log.error("Kit generation failed: %s", exc)
            raise HTTPException(status_code=503, detail="llm_unavailable") from exc
        await cache_set(redis, key, raw, get_settings().llm_cache_ttl_sec)
    try:
        data = extract_json_object(raw)
        item_payloads = _normalize_items(data, stages, global_complexity=body.global_complexity)
    except (ValueError, ValidationError, json.JSONDecodeError) as exc:
        log.error("Invalid kit JSON: %s", exc)
        repair_prompt = prompt + JSON_ONLY_SUFFIX
        try:
            raw2 = await llm.chat_completion(repair_prompt)
            item_payloads = _normalize_items(extract_json_object(raw2), stages, global_complexity=body.global_complexity)
        except Exception:
            item_payloads = _normalize_items({}, stages, global_complexity=body.global_complexity)
    kit = Kit(
        user_email=None,
        lesson_plan_id=plan.id if plan else None,
        name=body.topic or "Комплект раздаток",
        topic=body.topic,
        grade=body.grade,
        subject=body.subject,
        lesson_type=body.lesson_type,
        mode=body.mode,
    )
    db.add(kit)
    await db.flush()
    items = [KitItem(kit_id=kit.id, **payload) for payload in item_payloads]
    db.add_all(items)
    await db.commit()
    loaded = await _get_kit_or_404(db, kit.id)
    return GenerateFromPlanResponse(kit_id=loaded.id, items=[KitItemRead.model_validate(item) for item in loaded.items])


@router.post("/generate-from-scratch", response_model=GenerateFromPlanResponse)
async def generate_from_scratch(
    request: Request,
    body: GenerateFromScratchRequest,
    db: Annotated[AsyncSession, Depends(get_session)],
    llm: Annotated[GigaChatClient, Depends(get_llm)],
) -> GenerateFromPlanResponse:
    stages = _default_stages(body.lesson_type)
    plan_text = json.dumps([s.model_dump(mode="json") for s in stages], ensure_ascii=False)
    return await generate_from_plan(
        request=request,
        body=GenerateFromPlanRequest(
            lesson_plan_content=plan_text,
            stages=stages,
            topic=body.topic,
            grade=body.grade,
            subject=body.subject,
            lesson_type=body.lesson_type,
            mode="from_scratch",
            global_complexity=body.global_complexity,
            pedagogical=body.pedagogical,
        ),
        db=db,
        llm=llm,
    )


@router.get("/kits", response_model=list[KitRead])
async def list_kits(db: Annotated[AsyncSession, Depends(get_session)]) -> list[KitRead]:
    """Список всех комплектов (JSON). HTML-страница — на /teacher/kits."""
    res = await db.execute(select(Kit).options(selectinload(Kit.items)).order_by(Kit.created_at.desc()).limit(100))
    return [KitRead.model_validate(kit) for kit in res.scalars().all()]


@router.delete("/kits/{kit_id}")
async def delete_kit(kit_id: uuid.UUID, db: Annotated[AsyncSession, Depends(get_session)]) -> dict[str, bool]:
    kit = await _get_kit_or_404(db, kit_id)
    await db.delete(kit)
    await db.commit()
    return {"success": True}


@router.post("/kits/{kit_id}/copy", response_model=KitRead)
async def copy_kit(kit_id: uuid.UUID, db: Annotated[AsyncSession, Depends(get_session)]) -> KitRead:
    kit = await _get_kit_or_404(db, kit_id)
    clone = Kit(
        user_email=kit.user_email,
        lesson_plan_id=kit.lesson_plan_id,
        parent_kit_id=kit.id,
        name=f"{kit.name or kit.topic or 'Комплект'} — копия",
        topic=kit.topic,
        grade=kit.grade,
        subject=kit.subject,
        lesson_type=kit.lesson_type,
        mode=kit.mode,
        kit_type=kit.kit_type,
        version=kit.version,
    )
    db.add(clone)
    await db.flush()
    db.add_all([
        KitItem(
            kit_id=clone.id,
            stage_name=item.stage_name,
            type=item.type,
            title=item.title,
            content_levels=item.content_levels,
            complexity_level=item.complexity_level,
            complexity_distribution=item.complexity_distribution,
            sort_order=item.sort_order,
            teacher_notes=item.teacher_notes,
            answer_key=item.answer_key,
            example_mistakes=item.example_mistakes,
        )
        for item in kit.items
    ])
    await db.commit()
    return KitRead.model_validate(await _get_kit_or_404(db, clone.id))


@router.post("/kits/{kit_id}/template", response_model=KitRead)
async def save_kit_as_template(
    kit_id: uuid.UUID,
    body: KitTemplateRequest,
    db: Annotated[AsyncSession, Depends(get_session)],
) -> KitRead:
    kit = await _get_kit_or_404(db, kit_id)
    kit.is_template = True
    kit.template_name = body.template_name
    await db.commit()
    return KitRead.model_validate(await _get_kit_or_404(db, kit_id))


@router.get("/kits/{kit_id}", response_model=KitRead)
async def get_kit(kit_id: uuid.UUID, db: Annotated[AsyncSession, Depends(get_session)]) -> KitRead:
    return KitRead.model_validate(await _get_kit_or_404(db, kit_id))


@router.get("/kits/{kit_id}/items", response_model=list[KitItemRead])
async def get_kit_items(kit_id: uuid.UUID, db: Annotated[AsyncSession, Depends(get_session)]) -> list[KitItemRead]:
    kit = await _get_kit_or_404(db, kit_id)
    return [KitItemRead.model_validate(item) for item in kit.items]


@router.post("/kits/{kit_id}/items", response_model=KitItemRead)
async def create_kit_item(
    kit_id: uuid.UUID,
    body: KitItemCreate,
    db: Annotated[AsyncSession, Depends(get_session)],
) -> KitItemRead:
    await _get_kit_or_404(db, kit_id)
    item = KitItem(kit_id=kit_id, **body.model_dump(mode="json"))
    db.add(item)
    await db.commit()
    await db.refresh(item)
    return KitItemRead.model_validate(item)


@router.put("/kit-items/{item_id}", response_model=KitItemRead)
async def update_kit_item(
    item_id: uuid.UUID,
    body: KitItemUpdate,
    db: Annotated[AsyncSession, Depends(get_session)],
) -> KitItemRead:
    res = await db.execute(select(KitItem).where(KitItem.id == item_id))
    item = res.scalar_one_or_none()
    if item is None:
        raise HTTPException(status_code=404, detail="kit_item_not_found")
    updates = body.model_dump(exclude_unset=True, mode="json")
    for key, value in updates.items():
        setattr(item, key, value)
    await db.commit()
    await db.refresh(item)
    return KitItemRead.model_validate(item)


@router.delete("/kit-items/{item_id}")
async def delete_kit_item(item_id: uuid.UUID, db: Annotated[AsyncSession, Depends(get_session)]) -> dict[str, bool]:
    res = await db.execute(select(KitItem).where(KitItem.id == item_id))
    item = res.scalar_one_or_none()
    if item is None:
        raise HTTPException(status_code=404, detail="kit_item_not_found")
    await db.delete(item)
    await db.commit()
    return {"success": True}


@router.put("/kits/{kit_id}/reorder")
async def reorder_kit_items(
    kit_id: uuid.UUID,
    body: KitItemReorderRequest,
    db: Annotated[AsyncSession, Depends(get_session)],
) -> dict[str, Any]:
    """Drag-and-drop переупорядочивание элементов комплекта."""
    kit = await _get_kit_or_404(db, kit_id)
    item_by_id = {item.id: item for item in kit.items}
    if set(body.order) != set(item_by_id):
        raise HTTPException(status_code=400, detail="order_mismatch")
    for idx, item_id in enumerate(body.order):
        item_by_id[item_id].sort_order = idx
    await db.commit()
    return {"success": True, "order": [str(i) for i in body.order]}


@router.post("/kits/{kit_id}/ungeneratable")
async def kit_ungeneratable(
    kit_id: uuid.UUID,
    db: Annotated[AsyncSession, Depends(get_session)],
) -> dict[str, Any]:
    """Добавить «рукописный» (негенерабельный) элемент в комплект."""
    kit = await _get_kit_or_404(db, kit_id)
    sort_order = max((item.sort_order or 0) for item in kit.items) + 1 if kit.items else 0
    placeholder = ContentLevels(
        basic="✏️ Нарисуй схему ручкой на бумаге и сфотографируй.",
        medium="✏️ Нарисуй схему ручкой на бумаге и сфотографируй.",
        advanced="✏️ Нарисуй схему ручкой и сфотографируй. Подпиши элементы.",
    )
    item = KitItem(
        kit_id=kit.id,
        stage_name="Творческая часть",
        type="schema",
        title="🔒 Рукописный элемент",
        content_levels=placeholder.model_dump(mode="json"),
        complexity_level=2,
        sort_order=sort_order,
        teacher_notes="Этот элемент создаётся учеником вручную — не для ИИ-генерации.",
        answer_key=None,
        example_mistakes=[{"kind": "ungeneratable", "note": "manual_only"}],
    )
    db.add(item)
    await db.commit()
    await db.refresh(item)
    return {"success": True, "item_id": str(item.id)}


@router.post("/kits/{kit_id}/hot-mistakes")
async def kit_hot_mistakes(
    kit_id: uuid.UUID,
    request: Request,
    db: Annotated[AsyncSession, Depends(get_session)],
    llm: Annotated[GigaChatClient, Depends(get_llm)],
) -> dict[str, Any]:
    """Промпт №11 (US-12) — генерирует «Горячую десятку» ошибок для комплекта."""
    kit = await _get_kit_or_404(db, kit_id)
    redis = get_redis(request)
    prompt = prompt_hot_mistakes(topic=kit.topic or "урок", grade=kit.grade, subject=kit.subject, count=10)
    key = cache_key_for_prompt(prompt, prefix="hotmistakes")
    raw = await cache_get(redis, key)
    if raw is None:
        try:
            raw = await llm.chat_completion(prompt)
        except (GigaChatError, httpx.HTTPError) as exc:
            log.error("Hot mistakes generation failed: %s", exc)
            raise HTTPException(status_code=503, detail="llm_unavailable") from exc
        await cache_set(redis, key, raw, get_settings().llm_cache_ttl_sec)
    try:
        data = extract_json_object(raw)
    except (ValueError, json.JSONDecodeError):
        try:
            raw2 = await llm.chat_completion(prompt + JSON_ONLY_SUFFIX)
            data = extract_json_object(raw2)
        except Exception as exc:
            log.error("Hot mistakes repair failed: %s", exc)
            data = {"mistakes": []}
    return {"kit_id": str(kit.id), "mistakes": data.get("mistakes", [])}


@router.post("/kit-items/{item_id}/generate-example")
async def generate_item_example(
    item_id: uuid.UUID,
    request: Request,
    db: Annotated[AsyncSession, Depends(get_session)],
    llm: Annotated[GigaChatClient, Depends(get_llm)],
) -> dict[str, Any]:
    """Промпт №12 (US-16) — добавить пример с типичными ошибками к раздатке."""
    res = await db.execute(select(KitItem).where(KitItem.id == item_id))
    item = res.scalar_one_or_none()
    if item is None:
        raise HTTPException(status_code=404, detail="kit_item_not_found")
    kit_res = await db.execute(select(Kit).where(Kit.id == item.kit_id))
    kit = kit_res.scalar_one_or_none()
    grade = kit.grade if kit else 5
    content = " ".join((item.content_levels or {}).values()) if item.content_levels else item.title
    prompt = prompt_example_with_mistakes(title=item.title, content=content, grade=grade)
    redis = get_redis(request)
    key = cache_key_for_prompt(prompt, prefix="example_mistakes")
    raw = await cache_get(redis, key)
    if raw is None:
        try:
            raw = await llm.chat_completion(prompt)
        except (GigaChatError, httpx.HTTPError) as exc:
            log.error("Example mistakes generation failed: %s", exc)
            raise HTTPException(status_code=503, detail="llm_unavailable") from exc
        await cache_set(redis, key, raw, get_settings().llm_cache_ttl_sec)
    try:
        data = extract_json_object(raw)
    except (ValueError, json.JSONDecodeError):
        try:
            raw2 = await llm.chat_completion(prompt + JSON_ONLY_SUFFIX)
            data = extract_json_object(raw2)
        except Exception:
            data = {"example_mistakes": []}
    examples = data.get("example_mistakes") or []
    if examples:
        item.example_mistakes = list(item.example_mistakes or []) + examples
        await db.commit()
        await db.refresh(item)
    return {"success": True, "examples": examples}


@router.put("/kit-items/{item_id}/example", response_model=KitItemRead)
async def update_kit_item_example(
    item_id: uuid.UUID,
    body: ExampleMistakesPayload,
    db: Annotated[AsyncSession, Depends(get_session)],
) -> KitItemRead:
    res = await db.execute(select(KitItem).where(KitItem.id == item_id))
    item = res.scalar_one_or_none()
    if item is None:
        raise HTTPException(status_code=404, detail="kit_item_not_found")
    item.example_mistakes = body.example_mistakes
    await db.commit()
    await db.refresh(item)
    return KitItemRead.model_validate(item)


@router.post("/kits/{kit_id}/variants")
async def create_kit_variants(kit_id: uuid.UUID, db: Annotated[AsyncSession, Depends(get_session)]) -> dict[str, Any]:
    base = await _get_kit_or_404(db, kit_id)
    config = {
        "A": {"name": "Слабые", "complexity": 1, "suffix": "вариант А"},
        "B": {"name": "Средние", "complexity": 2, "suffix": "вариант Б"},
        "C": {"name": "Сильные", "complexity": 3, "suffix": "вариант В"},
    }
    created: list[KitRead] = []
    for code, cfg in config.items():
        variant = Kit(
            user_email=base.user_email,
            lesson_plan_id=base.lesson_plan_id,
            parent_kit_id=base.id,
            name=f"{base.name or base.topic or 'Комплект'} — {cfg['suffix']}",
            topic=base.topic,
            grade=base.grade,
            subject=base.subject,
            lesson_type=base.lesson_type,
            mode="variant",
            kit_type=code,
            version=base.version,
        )
        db.add(variant)
        await db.flush()
        db.add_all([
            KitItem(
                kit_id=variant.id,
                stage_name=item.stage_name,
                type=item.type,
                title=f"{item.title} ({cfg['name']})",
                content_levels=item.content_levels,
                complexity_level=cfg["complexity"],
                complexity_distribution=item.complexity_distribution,
                sort_order=item.sort_order,
                teacher_notes=item.teacher_notes,
                answer_key=item.answer_key,
                example_mistakes=item.example_mistakes,
            )
            for item in base.items
        ])
        await db.flush()
        created.append(KitRead.model_validate(await _get_kit_or_404(db, variant.id)))
    await db.commit()
    return {"variants": created}


@router.post("/kits/{kit_id}/fork", response_model=KitRead)
async def fork_kit(
    kit_id: uuid.UUID,
    body: LessonFeedback,
    db: Annotated[AsyncSession, Depends(get_session)],
) -> KitRead:
    kit = await _get_kit_or_404(db, kit_id)
    fork = Kit(
        user_email=kit.user_email,
        lesson_plan_id=kit.lesson_plan_id,
        parent_kit_id=kit.id,
        name=f"{kit.name or kit.topic or 'Комплект'} — v{kit.version + 1}",
        topic=kit.topic,
        grade=kit.grade,
        subject=kit.subject,
        lesson_type=kit.lesson_type,
        mode="fork",
        kit_type=kit.kit_type,
        version=kit.version + 1,
        feedback=body.model_dump(mode="json"),
    )
    db.add(fork)
    await db.flush()
    db.add_all([
        KitItem(
            kit_id=fork.id,
            stage_name=item.stage_name,
            type=item.type,
            title=item.title,
            content_levels=item.content_levels,
            complexity_level=item.complexity_level,
            complexity_distribution=item.complexity_distribution,
            sort_order=item.sort_order,
            teacher_notes=item.teacher_notes,
            answer_key=item.answer_key,
            example_mistakes=item.example_mistakes,
        )
        for item in kit.items
    ])
    await db.commit()
    return KitRead.model_validate(await _get_kit_or_404(db, fork.id))


@router.get("/design-settings", response_model=DesignSettingsPayload)
async def get_design_settings(db: Annotated[AsyncSession, Depends(get_session)]) -> DesignSettingsPayload:
    res = await db.execute(select(DesignSettings).where(DesignSettings.user_email == DEMO_EMAIL))
    settings = res.scalar_one_or_none()
    if settings is None:
        return DesignSettingsPayload()
    return DesignSettingsPayload.model_validate(settings, from_attributes=True)


@router.put("/design-settings", response_model=DesignSettingsPayload)
async def put_design_settings(
    body: DesignSettingsPayload,
    db: Annotated[AsyncSession, Depends(get_session)],
) -> DesignSettingsPayload:
    user_res = await db.execute(select(User).where(User.email == DEMO_EMAIL))
    if user_res.scalar_one_or_none() is None:
        db.add(User(email=DEMO_EMAIL, password_hash=None, demo_id="demo"))
        await db.flush()
    res = await db.execute(select(DesignSettings).where(DesignSettings.user_email == DEMO_EMAIL))
    settings = res.scalar_one_or_none()
    if settings is None:
        settings = DesignSettings(user_email=DEMO_EMAIL, **body.model_dump(mode="json"))
        db.add(settings)
    else:
        for key, value in body.model_dump(mode="json").items():
            setattr(settings, key, value)
    await db.commit()
    return body


@router.get("/lesson-plans", response_model=list[LessonPlanRead])
async def list_lesson_plans(db: Annotated[AsyncSession, Depends(get_session)]) -> list[LessonPlanRead]:
    res = await db.execute(select(LessonPlan).order_by(LessonPlan.created_at.desc()).limit(50))
    return [LessonPlanRead.model_validate(plan) for plan in res.scalars().all()]


@router.delete("/lesson-plans/{lesson_plan_id}")
async def delete_lesson_plan(lesson_plan_id: uuid.UUID, db: Annotated[AsyncSession, Depends(get_session)]) -> dict[str, bool]:
    res = await db.execute(select(LessonPlan).where(LessonPlan.id == lesson_plan_id))
    plan = res.scalar_one_or_none()
    if plan is None:
        raise HTTPException(status_code=404, detail="lesson_plan_not_found")
    await db.delete(plan)
    await db.commit()
    return {"success": True}


@router.post("/export/kit/{kit_id}")
async def export_kit_pdf(
    kit_id: uuid.UUID,
    body: KitExportRequest,
    db: Annotated[AsyncSession, Depends(get_session)],
) -> Response:
    try:
        pdf = await render_kit_pdf(kit_id, db, export=body)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    return Response(
        content=pdf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="beresta_kit_{kit_id}.pdf"'},
    )


@router.post("/export/kit/{kit_id}/pdf")
async def export_kit_pdf_alias(
    kit_id: uuid.UUID,
    body: KitExportRequest,
    db: Annotated[AsyncSession, Depends(get_session)],
) -> Response:
    return await export_kit_pdf(kit_id=kit_id, body=body, db=db)


@router.post("/export/kit/{kit_id}/docx")
async def export_kit_docx(kit_id: uuid.UUID, db: Annotated[AsyncSession, Depends(get_session)]) -> Response:
    """DOCX-экспорт с настоящими таблицами Word под каждый уровень."""
    kit = await _get_kit_or_404(db, kit_id)
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="python_docx_not_installed") from exc

    document = Document()
    document.add_heading(kit.name or kit.topic or "Комплект Beresta", level=1)
    meta = document.add_paragraph()
    meta.add_run(f"{kit.subject}, {kit.grade} класс — {len(kit.items)} раздаток").italic = True

    sorted_items = sorted(kit.items, key=lambda i: (i.sort_order or 0, i.created_at))
    for idx, item in enumerate(sorted_items, start=1):
        document.add_heading(f"{idx}. {item.title}", level=2)
        document.add_paragraph(f"Этап: {item.stage_name or '—'} · Тип: {item.type}")

        table = document.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Уровень"
        hdr[1].text = "Содержание"
        for level, label in (("basic", "🟢 Базовый"), ("medium", "🟡 Средний"), ("advanced", "🔴 Продвинутый")):
            row = table.add_row().cells
            row[0].text = label
            row[1].text = (item.content_levels or {}).get(level, "")

        if item.answer_key:
            document.add_paragraph().add_run("Ответы:").bold = True
            ak_table = document.add_table(rows=1, cols=2)
            ak_table.style = "Light Grid"
            hk = ak_table.rows[0].cells
            hk[0].text = "Уровень"
            hk[1].text = "Ответ"
            for k, v in (item.answer_key or {}).items():
                row = ak_table.add_row().cells
                row[0].text = str(k)
                row[1].text = str(v)
        if item.teacher_notes:
            document.add_paragraph().add_run("Методические заметки:").bold = True
            document.add_paragraph(item.teacher_notes)
        if item.example_mistakes:
            document.add_paragraph().add_run("Примеры ошибок:").bold = True
            for ex in item.example_mistakes:
                if isinstance(ex, dict):
                    parts = [f"{k}: {v}" for k, v in ex.items()]
                    document.add_paragraph(" · ".join(parts), style="List Bullet")
        document.add_paragraph().add_run("")  # spacer
    buf = io.BytesIO()
    document.save(buf)
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="beresta_kit_{kit_id}.docx"'},
    )


@router.post("/export/kit/{kit_id}/zip")
async def export_kit_zip(
    kit_id: uuid.UUID,
    db: Annotated[AsyncSession, Depends(get_session)],
) -> Response:
    """ZIP с папками по этапам + PDF (если доступен WeasyPrint) + answers JSON."""
    import zipfile

    kit = await _get_kit_or_404(db, kit_id)
    try:
        pdf_bytes = await render_kit_pdf(kit_id, db, export=KitExportRequest())
    except Exception as exc:
        log.warning("ZIP export: kit PDF unavailable: %s", exc)
        pdf_bytes = None
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        root = f"{kit.topic or 'urok'}_{kit.grade}_klass"
        zf.writestr(
            f"{root}/README.txt",
            f"Комплект Beresta: {kit.name or kit.topic}\n"
            f"Предмет: {kit.subject}, {kit.grade} класс\n"
            f"Этапов с раздатками: {len(kit.items)}\n",
        )
        if pdf_bytes:
            zf.writestr(f"{root}/КОМПЛЕКТ_УЧЕНИКА.pdf", pdf_bytes)
            try:
                pdf_teacher = await render_kit_pdf(
                    kit_id, db, export=KitExportRequest(version="teacher", title_page=True)
                )
                zf.writestr(f"{root}/учительская_версия/КОМПЛЕКТ_УЧИТЕЛЯ.pdf", pdf_teacher)
            except Exception as exc:
                log.warning("Teacher PDF unavailable: %s", exc)
        sorted_items = sorted(kit.items, key=lambda i: (i.sort_order or 0, i.created_at))
        for idx, item in enumerate(sorted_items, start=1):
            folder = f"{root}/{idx:02d}_{item.stage_name or 'stage'}"
            for level, label in (("basic", "слабые"), ("medium", "средние"), ("advanced", "сильные")):
                zf.writestr(
                    f"{folder}/{label}_{item.type}.txt",
                    f"{item.title}\n{'=' * len(item.title)}\n\n{(item.content_levels or {}).get(level, '')}\n",
                )
            if item.answer_key or item.teacher_notes or item.example_mistakes:
                zf.writestr(
                    f"{root}/учительская_версия/{idx:02d}_{item.type}_ответы.json",
                    json.dumps(
                        {
                            "title": item.title,
                            "answer_key": item.answer_key,
                            "teacher_notes": item.teacher_notes,
                            "example_mistakes": item.example_mistakes,
                        },
                        ensure_ascii=False,
                        indent=2,
                    ),
                )
    return Response(
        content=buf.getvalue(),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="beresta_kit_{kit_id}.zip"'},
    )
